from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, BackgroundTasks
from sqlalchemy.ext.asyncio import AsyncSession
from uuid import UUID
from typing import List
import logging

from app.core.database import get_db, AsyncSessionLocal
from app.core.config import settings
from app.core.security import get_current_user_id
from app.schemas.chat_schema import (
    AIChatSessionCreate,
    AIChatSessionResponse,
    AIChatSessionWithMessagesResponse,
    AIChatMessageCreate,
    AIChatMessageResponse,
    AIChatSessionUpdate
)
from app.services.chat_service import ChatService
from app.services.assistant_service import ai_assistant_service
from app.models.chat import AIChatMessage

router = APIRouter()

@router.post("/sessions", response_model=AIChatSessionResponse, status_code=status.HTTP_201_CREATED)
async def create_chat_session(
    payload: AIChatSessionCreate,
    db: AsyncSession = Depends(get_db),
    user_id: UUID = Depends(get_current_user_id)
):
    """
    Tạo mới một phiên trò chuyện AI trong một dự án.
    Bắt buộc người dùng phải có quyền truy cập dự án.
    """
    session = await ChatService.create_session(
        db=db,
        project_id=payload.project_id,
        user_id=user_id,
        title=payload.title
    )
    return session

@router.get("/sessions", response_model=List[AIChatSessionResponse])
async def get_chat_sessions(
    project_id: UUID,
    db: AsyncSession = Depends(get_db),
    user_id: UUID = Depends(get_current_user_id)
):
    """
    Lấy danh sách các phiên trò chuyện của người dùng trong một dự án.
    """
    sessions = await ChatService.get_sessions(db=db, project_id=project_id, user_id=user_id)
    return sessions

@router.get("/sessions/{session_id}/messages", response_model=AIChatSessionWithMessagesResponse)
async def get_session_history(
    session_id: UUID,
    db: AsyncSession = Depends(get_db),
    user_id: UUID = Depends(get_current_user_id)
):
    """
    Lấy chi tiết phiên trò chuyện kèm theo toàn bộ lịch sử tin nhắn.
    """
    session = await ChatService.get_session_with_messages(db=db, session_id=session_id, user_id=user_id)
    return session

logger = logging.getLogger(__name__)

async def run_assistant_background(session_id: UUID, user_id: UUID, new_prompt_content: str):
    logger.info(f"Starting background AI assistant session {session_id} for user {user_id}")
    async with AsyncSessionLocal() as db:
        try:
            # 1. Load session with history
            session = await ChatService.get_session_with_messages(db=db, session_id=session_id, user_id=user_id)
            history = session.messages
            
            # The last message is the user prompt we just saved in the route.
            # We pass history[:-1] to chat_with_assistant to avoid duplicating the prompt in Gemini history.
            # To optimize token usage and avoid rate limits, we limit the active context to the last 20 messages (sliding window).
            history_before_prompt = history[:-1] if history else []
            history_before_prompt = history_before_prompt[-20:]

            # Helper to save messages using its own session to avoid cross-coroutine session conflicts
            async def save_message_callback(
                role: str,
                content: str = None,
                tool_calls: list = None,
                tool_results: list = None,
                thought_signature: str = None
            ):
                if role == "user":
                    # The user message was already saved synchronously in the route
                    return None
                async with AsyncSessionLocal() as write_db:
                    msg = await ChatService.create_message(
                        db=write_db,
                        session_id=session_id,
                        role=role,
                        content=content,
                        tool_calls=tool_calls,
                        tool_results=tool_results,
                        thought_signature=thought_signature
                    )
                    await write_db.commit()
                    return msg
            
            # 2. Run assistant service
            await ai_assistant_service.chat_with_assistant(
                user_id=user_id,
                project_id=session.project_id,
                history=history_before_prompt,
                new_prompt=new_prompt_content,
                message_saver_callback=save_message_callback
            )
            
        except Exception as e:
            logger.exception(f"Error in run_assistant_background for session {session_id}: {e}")
            try:
                async with AsyncSessionLocal() as read_db:
                    session = await ChatService.get_session_with_messages(db=read_db, session_id=session_id, user_id=user_id)
                    has_tools = any(m.tool_calls for m in session.messages) if session else False
                
                async with AsyncSessionLocal() as err_db:
                    if has_tools:
                        content = "⏳ Quá trình có thể lâu hơn chút do quá nhiều yêu cầu."
                    else:
                        content = "⏳ Trợ lý AI đang gặp tình trạng quá tải. Vui lòng thử lại sau 1 phút nhé."
                    
                    await ChatService.create_message(
                        db=err_db,
                        session_id=session_id,
                        role="assistant",
                        content=content,
                        tool_calls=None,
                        tool_results=None
                    )
                    await err_db.commit()
            except Exception as se:
                logger.error(f"Failed to write error message to DB: {se}")

@router.post("/sessions/{session_id}/messages", response_model=AIChatMessageResponse)
async def send_chat_message(
    session_id: UUID,
    payload: AIChatMessageCreate,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    user_id: UUID = Depends(get_current_user_id)
):
    """
    Gửi câu hỏi của người dùng và nhận câu trả lời từ Trợ lý AI.
    Hệ thống sẽ chạy luồng suy luận của Gemini dưới nền và phản hồi ngay lập tức.
    """
    # 1. Verify access and load session
    session = await ChatService.get_session_with_messages(db=db, session_id=session_id, user_id=user_id)
    
    # 2. Save user message to database immediately
    user_msg = await ChatService.create_message(
        db=db,
        session_id=session_id,
        role="user",
        content=payload.content
    )
    
    # 3. Add background task to run the AI assistant
    background_tasks.add_task(
        run_assistant_background,
        session_id=session_id,
        user_id=user_id,
        new_prompt_content=payload.content
    )
    
    return user_msg

@router.post("/sessions/{session_id}/stop", response_model=AIChatMessageResponse)
async def stop_chat_session(
    session_id: UUID,
    db: AsyncSession = Depends(get_db),
    user_id: UUID = Depends(get_current_user_id)
):
    """
    Dừng tiến trình chạy của Trợ lý AI và ghi nhận tin nhắn dừng.
    """
    session = await ChatService.get_session_with_messages(db=db, session_id=session_id, user_id=user_id)
    
    msg = await ChatService.create_message(
        db=db,
        session_id=session_id,
        role="assistant",
        content="⏹️ Tiến trình tạo công việc đã bị dừng bởi người dùng.",
        tool_calls=None,
        tool_results=None
    )
    return msg

@router.patch("/sessions/{session_id}", response_model=AIChatSessionResponse)
async def rename_chat_session(
    session_id: UUID,
    payload: AIChatSessionUpdate,
    db: AsyncSession = Depends(get_db),
    user_id: UUID = Depends(get_current_user_id)
):
    """
    Đổi tên một phiên trò chuyện AI.
    """
    return await ChatService.rename_session(
        db=db,
        session_id=session_id,
        user_id=user_id,
        title=payload.title
    )

@router.delete("/sessions/{session_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_chat_session(
    session_id: UUID,
    db: AsyncSession = Depends(get_db),
    user_id: UUID = Depends(get_current_user_id)
):
    """
    Xóa một phiên trò chuyện AI và toàn bộ lịch sử tin nhắn của nó.
    """
    await ChatService.delete_session(db=db, session_id=session_id, user_id=user_id)
    return None

async def extract_text_from_file(file_name: str, content: bytes) -> str:
    ext = file_name.split(".")[-1].lower()
    if ext == "pdf":
        try:
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(content))
            text = []
            for page in reader.pages:
                text.append(page.extract_text() or "")
            return "\n".join(text)
        except Exception as e:
            raise HTTPException(
                status_code=400,
                detail=f"Không thể đọc file PDF: {str(e)}"
            )
    elif ext == "docx":
        try:
            import docx
            import io
            doc = docx.Document(io.BytesIO(content))
            text = []
            for para in doc.paragraphs:
                text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text.append(" | ".join(row_text))
            return "\n".join(text)
        except Exception as e:
            raise HTTPException(
                status_code=400,
                detail=f"Không thể đọc file Word (.docx): {str(e)}"
            )
    elif ext == "xlsx":
        try:
            import openpyxl
            import io
            wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            text = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text.append(f"--- Sheet: {sheet_name} ---")
                for row in sheet.iter_rows(values_only=True):
                    if any(val is not None for val in row):
                        row_vals = [str(val) if val is not None else "" for val in row]
                        text.append(" | ".join(row_vals))
            return "\n".join(text)
        except Exception as e:
            raise HTTPException(
                status_code=400,
                detail=f"Không thể đọc file Excel (.xlsx): {str(e)}"
            )
    else:
        # Treat as plain text
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                return content.decode("latin-1")
            except Exception as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Không thể đọc định dạng file văn bản: {str(e)}"
                )

@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    user_id: UUID = Depends(get_current_user_id)
):
    """
    Tải lên tài liệu (PDF, TXT, MD, v.v.), trích xuất văn bản và trả về.
    Giới hạn kích thước tệp là 2MB.
    """
    MAX_FILE_SIZE = 2 * 1024 * 1024  # 2MB
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Kích thước tệp vượt quá giới hạn cho phép (Tối đa 2MB)."
        )
    
    extracted_text = await extract_text_from_file(file.filename, content)
    estimated_tokens = int(len(extracted_text) * 0.35)
    
    return {
        "fileName": file.filename,
        "fileSize": f"{len(content) // 1024} KB",
        "content": extracted_text,
        "estimatedTokens": estimated_tokens
    }

from pydantic import BaseModel

class ExtractProjectDocumentRequest(BaseModel):
    fileUrl: str
    fileName: str

@router.post("/extract-project-document")
async def extract_project_document(
    request_data: ExtractProjectDocumentRequest,
    user_id: UUID = Depends(get_current_user_id)
):
    """
    Tải tệp tin từ tài liệu dự án, trích xuất văn bản và trả về cho AI.
    """
    import httpx
    
    file_url = request_data.fileUrl
    if not file_url.startswith("http"):
        file_url = f"{settings.PM_SERVICE_BASE_URL}{file_url}"
        
    try:
        async with httpx.AsyncClient(timeout=15.0) as client:
            response = await client.get(file_url)
            if response.status_code != 200:
                raise HTTPException(
                    status_code=400,
                    detail=f"Không thể tải tệp tin từ tài liệu dự án (Mã lỗi: {response.status_code})"
                )
            content = response.content
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Không thể kết nối tải tệp tin từ tài liệu dự án: {str(e)}"
        )
        
    MAX_FILE_SIZE = 10 * 1024 * 1024  # Giới hạn 10MB
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail="Tệp tin vượt quá kích thước xử lý tối đa (10MB)."
        )
        
    extracted_text = await extract_text_from_file(request_data.fileName, content)
    estimated_tokens = int(len(extracted_text) * 0.35)
    
    file_size_kb = len(content) // 1024
    file_size_str = f"{file_size_kb} KB"
    if file_size_kb > 1024:
        file_size_str = f"{file_size_kb / 1024:.1f} MB"
        
    return {
        "fileName": request_data.fileName,
        "fileSize": file_size_str,
        "content": extracted_text,
        "estimatedTokens": estimated_tokens
    }
